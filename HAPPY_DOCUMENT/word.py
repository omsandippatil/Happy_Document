import comtypes.client
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, Listbox, simpledialog
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import os
from docx import Document
import pandas as pd
import time

# Function to extract text from a Word document (.docx and .doc)
def extract_text_word(file_path):
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif file_path.endswith('.doc'):
        text = ""
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(file_path))
        for para in doc.Paragraphs:
            text += para.Range.Text + "\n"
        doc.Close()
        word.Quit()
    else:
        raise ValueError("Unsupported file format.")
    
    return text

# Function to merge multiple Word files into one
def merge_word_files(file_paths, output_path, progress_var):
    merged_doc = Document()
    for i, file in enumerate(file_paths):
        if file.endswith('.docx'):
            doc = Document(file)
        elif file.endswith('.doc'):
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(file))
        else:
            raise ValueError("Unsupported file format.")
        
        for para in doc.paragraphs:
            merged_doc.add_paragraph(para.text)
        
        if file.endswith('.doc'):
            doc.Close()
            word.Quit()

        # Update progress bar
        progress_var.set((i + 1) / len(file_paths) * 100)
        time.sleep(0.5)  # Simulated processing time

    merged_doc.save(output_path)

# Function to convert Word to PDF using Microsoft Word via COM
def word_to_pdf(word_file, output_pdf, progress_var):
    try:
        if not os.path.exists(word_file):
            raise FileNotFoundError(f"The file {word_file} does not exist.")

        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(word_file))
        doc.SaveAs(os.path.abspath(output_pdf), FileFormat=17)  # 17 is the format ID for PDF in Word
        doc.Close()
        word.Quit()

        print(f"Converted {word_file} to {output_pdf} successfully.")
    except Exception as e:
        print(f"An error occurred while converting to PDF: {str(e)}")

# Function to search text in a Word document
def search_in_word(word_file, search_text):
    if word_file.endswith('.docx'):
        doc = Document(word_file)
    elif word_file.endswith('.doc'):
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(word_file))
    else:
        raise ValueError("Unsupported file format.")

    results = []
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            results.append((i, para.text))  # Store index and text

    if word_file.endswith('.doc'):
        doc.Close()
        word.Quit()

    return results

# Function to set password protection on a Word document using COM
def set_word_password(word_file, output_file, password):
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        
        # Open the document
        doc = word.Documents.Open(os.path.abspath(word_file))

        # Set the password and save the document
        doc.SaveAs(os.path.abspath(output_file), Password=password)
        
        # Close the document and quit Word
        doc.Close()
        word.Quit()
        
        print(f"Password set successfully for {output_file}.")
    except Exception as e:
        print(f"An error occurred: {str(e)}")

class WordUtilityGUI:
    def __init__(self, master):
        self.master = master
        master.title("Word Utility")
        master.geometry("900x600")
        master.resizable(True, True)

        # Create a style
        self.style = ttk.Style(theme="darkly")

        # Main frame
        self.main_frame = ttk.Frame(master, padding="20 20 20 20")
        self.main_frame.pack(fill=BOTH, expand=YES)

        # Title
        ttk.Label(self.main_frame, text="Word Utility", font=("Helvetica", 24, "bold")).pack(pady=20)

        # Operation selection frame
        self.operation_frame = ttk.LabelFrame(self.main_frame, text="Select Operation", padding="20 20 20 20")
        self.operation_frame.pack(fill=X, padx=10, pady=10)

        self.operation = tk.StringVar(value="extract")
        operations = [
            ("Extract Text", "extract"),
            ("Merge Word Files", "merge"),
            ("Convert to PDF", "convert_pdf"),
            ("Convert to Excel", "convert_excel"),
            ("Set Password", "set_password"),
            ("Search Text", "search_text"),
        ]

        for text, value in operations:
            ttk.Radiobutton(self.operation_frame, text=text, variable=self.operation, value=value).pack(side=LEFT, padx=10)

        # File selection frame
        self.file_frame = ttk.Frame(self.main_frame)
        self.file_frame.pack(fill=X, padx=10, pady=20)

        self.file_button = ttk.Button(self.file_frame, text="Select Word File(s)", command=self.browse_files, width=30)
        self.file_button.pack(side=LEFT, padx=30)

        self.file_label = ttk.Label(self.file_frame, text="No files selected")
        self.file_label.pack(side=LEFT, padx=10)

        # Process button
        self.process_button = ttk.Button(self.main_frame, text="Process", command=self.process, style="success.TButton", width=40)
        self.process_button.pack(pady=40)

        # Status bar
        self.status_var = tk.StringVar()
        self.status_bar = ttk.Label(self.main_frame, textvariable=self.status_var, relief=SUNKEN, anchor=W)
        self.status_bar.pack(fill=X, side=BOTTOM, pady=15)

        # Progress bar
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(self.main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=X, padx=10, pady=10)

        self.selected_files = []

    def browse_files(self):
        if self.operation.get() in ["extract", "convert_pdf", "convert_excel", "search_text"]:
            files = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx;*.doc")])
        else:
            files = filedialog.askopenfilenames(filetypes=[("Word Files", "*.docx;*.doc")])

        if isinstance(files, str):  # Single file for extract, convert to PDF, Excel, or search text
            self.selected_files = [files] if files else []
        else:  # Multiple files for merge
            self.selected_files = list(files)

        num_files = len(self.selected_files)
        self.file_label.config(text=f"{num_files} file{'s' if num_files != 1 else ''} selected")

        if self.operation.get() == "merge":
            self.show_sorting_window()

    def show_sorting_window(self):
        sorting_window = tk.Toplevel(self.master)
        sorting_window.title("Sort Selected Files")
        sorting_window.geometry("400x300")

        self.sorting_listbox = Listbox(sorting_window, selectmode=tk.SINGLE, height=10)
        self.sorting_listbox.pack(fill=BOTH, expand=True)

        for file in self.selected_files:
            self.sorting_listbox.insert(tk.END, os.path.basename(file))

        self.sorting_listbox.bind("<Button-1>", self.on_click)
        self.sorting_listbox.bind("<B1-Motion>", self.on_drag)

        # Note label for drag-and-drop functionality
        note_label = ttk.Label(sorting_window, text="Drag and drop files here to reorder.", font=("Helvetica", 10, "italic"))
        note_label.pack(pady=5)

        ttk.Button(sorting_window, text="Done", command=lambda: self.finish_sorting(sorting_window)).pack(side=tk.RIGHT, padx=10, pady=10)

        self.dragged_index = None

    def on_click(self, event):
        self.dragged_index = self.sorting_listbox.nearest(event.y)

    def on_drag(self, event):
        if self.dragged_index is None:
            return
        current_index = self.sorting_listbox.nearest(event.y)
        if current_index != self.dragged_index:
            item_text = self.sorting_listbox.get(self.dragged_index)
            self.sorting_listbox.delete(self.dragged_index)
            self.sorting_listbox.insert(current_index, item_text)
            self.dragged_index = current_index

    def finish_sorting(self, window):
        self.selected_files = [self.selected_files[int(i)] for i in range(len(self.sorting_listbox.get(0, tk.END))) if i < len(self.selected_files)]
        window.destroy()

    def process(self):
        operation = self.operation.get()

        if operation == "extract":
            if not self.selected_files:
                messagebox.showerror("Error", "Please select a Word file to extract text from.")
                return
            
            text = extract_text_word(self.selected_files[0])
            self.show_text_window(text)
        elif operation == "merge":
            if not self.selected_files:
                messagebox.showerror("Error", "Please select Word files to merge.")
                return

            output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
            if not output_file:
                return
            
            self.progress_var.set(0)
            self.status_var.set("Merging files...")
            self.master.update()
            merge_word_files(self.selected_files, output_file, self.progress_var)
            messagebox.showinfo("Success", "Files merged successfully.")
            self.status_var.set("Merging complete.")
        elif operation == "convert_pdf":
            if not self.selected_files:
                messagebox.showerror("Error", "Please select a Word file to convert to PDF.")
                return

            output_pdf = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
            if not output_pdf:
                return
            
            self.progress_var.set(0)
            self.status_var.set("Converting to PDF...")
            self.master.update()
            word_to_pdf(self.selected_files[0], output_pdf, self.progress_var)
            messagebox.showinfo("Success", "File converted to PDF successfully.")
            self.status_var.set("Conversion complete.")
        elif operation == "set_password":
            if not self.selected_files:
                messagebox.showerror("Error", "Please select a Word file to set password.")
                return
            
            password = simpledialog.askstring("Password", "Enter password for the document:", show='*')
            if password is None:
                return
            
            output_file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
            if not output_file:
                return
            
            self.progress_var.set(0)
            self.status_var.set("Setting password...")
            self.master.update()
            set_word_password(self.selected_files[0], output_file, password)
            messagebox.showinfo("Success", "Password set successfully.")
            self.status_var.set("Password set.")
        elif operation == "search_text":
            if not self.selected_files:
                messagebox.showerror("Error", "Please select a Word file to search text.")
                return
            
            search_text = simpledialog.askstring("Search", "Enter text to search:")
            if search_text is None or search_text.strip() == "":
                return
            
            self.progress_var.set(0)
            self.status_var.set("Searching text...")
            self.master.update()
            results = search_in_word(self.selected_files[0], search_text)
            if results:
                result_text = "\n".join([f"Line {i+1}: {text}" for i, text in results])
                self.show_text_window(result_text)
            else:
                messagebox.showinfo("Result", "No matches found.")
            self.status_var.set("Search complete.")

    def show_text_window(self, text):
        text_window = tk.Toplevel(self.master)
        text_window.title("Extracted Text")
        text_window.geometry("600x400")

        text_area = scrolledtext.ScrolledText(text_window, wrap=tk.WORD, font=("Helvetica", 12))
        text_area.pack(fill=BOTH, expand=True)
        text_area.insert(tk.END, text)
        text_area.config(state=tk.DISABLED)

        ttk.Button(text_window, text="Close", command=text_window.destroy).pack(pady=10)

if __name__ == "__main__":
    root = tk.Tk()
    app = WordUtilityGUI(root)
    root.mainloop()
