import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
from openpyxl import load_workbook
from fpdf import FPDF
from docx import Document
import os
import win32com.client as win32
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import seaborn as sns

class EnhancedExcelUtilityApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Enhanced Excel Utility")
        self.geometry("1000x700")
        self.configure(bg='#2E2E2E')
        
        # Configure modern styling
        self.style = ttk.Style(self)
        self.style.theme_use('clam')
        self.configure_styles()
        
        # Main container
        self.main_container = ttk.Frame(self)
        self.main_container.pack(fill='both', expand=True, padx=10, pady=10)
        
        # File selection frame
        self.create_file_frame()
        
        # Notebook for tabs
        self.notebook = ttk.Notebook(self.main_container)
        self.notebook.pack(fill='both', expand=True, pady=(10, 0))
        
        # Initialize variables
        self.selected_file = ""
        self.df = None
        
        # Create tabs
        self.create_tabs()
        
        # Status bar
        self.status_var = tk.StringVar()
        self.status_bar = ttk.Label(self, textvariable=self.status_var, style='Status.TLabel')
        self.status_bar.pack(fill='x', pady=5)
        self.status_var.set("Ready")

    def configure_styles(self):
        # Configure modern styles
        self.style.configure("TFrame", background='#2E2E2E')
        self.style.configure("TLabel", background='#2E2E2E', foreground='white', font=('Helvetica', 16))
        self.style.configure("Header.TLabel", font=('Helvetica', 16, 'bold'))
        self.style.configure("Status.TLabel", background='#2c2c2c', foreground='#00FF00', font=('Helvetica', 10))
        
        # Button styling
        self.style.configure("TButton",
                           font=('Helvetica', 12),
                           padding=10,
                           background='#4A4A4A',
                           foreground='black')
        self.style.map("TButton",
                      background=[('active', '#5A5A5A')],
                      foreground=[('active', 'white')])
        
        # Notebook styling
        self.style.configure("TNotebook", background='#2E2E2E', padding=5)
        self.style.configure("TNotebook.Tab", padding=[12, 4],
                           font=('Helvetica', 12),
                           background='#4A4A4A',
                           foreground='black')

    def create_file_frame(self):
        file_frame = ttk.Frame(self.main_container)
        file_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(file_frame, text="Current File:", style='Header.TLabel').pack(side='left', padx=5)
        self.file_label = ttk.Label(file_frame, text="No file selected")
        self.file_label.pack(side='left', padx=5)
        
        ttk.Button(file_frame, text="Browse", command=self.browse_file).pack(side='right', padx=5)

    def create_tabs(self):
        # Create and add all tabs
        self.tabs = {
            "Convert to CSV": self.create_csv_tab,
            "Filter & Sort": self.create_filter_sort_tab,
            "Pivot Tables": self.create_pivot_tab,
            "Data Visualization": self.create_visualization_tab,
            "Convert to Word": self.create_word_tab
        }
        
        for tab_name, tab_function in self.tabs.items():
            frame = ttk.Frame(self.notebook)
            self.notebook.add(frame, text=tab_name)
            tab_function(frame)

    def create_csv_tab(self, frame):
        ttk.Label(frame, text="Convert Excel to CSV", style='Header.TLabel').pack(pady=10)
        
        options_frame = ttk.Frame(frame)
        options_frame.pack(fill='x', padx=20, pady=10)
        
        ttk.Label(options_frame, text="Sheet Name:").pack(side='left', padx=5)
        self.csv_sheet_var = tk.StringVar()
        self.csv_sheet_combo = ttk.Combobox(options_frame, textvariable=self.csv_sheet_var)
        self.csv_sheet_combo.pack(side='left', padx=5)
        
        ttk.Button(frame, text="Convert to CSV", command=self.convert_to_csv).pack(pady=10)

    def create_filter_sort_tab(self, frame):
        ttk.Label(frame, text="Filter and Sort Data", style='Header.TLabel').pack(pady=10)
        
        # Column selection
        col_frame = ttk.Frame(frame)
        col_frame.pack(fill='x', padx=20, pady=10)
        
        ttk.Label(col_frame, text="Column:").pack(side='left', padx=5)
        self.filter_column_var = tk.StringVar()
        self.filter_column_combo = ttk.Combobox(col_frame, textvariable=self.filter_column_var)
        self.filter_column_combo.pack(side='left', padx=5)
        
        # Filter value
        filter_frame = ttk.Frame(frame)
        filter_frame.pack(fill='x', padx=20, pady=10)
        
        ttk.Label(filter_frame, text="Filter Value:").pack(side='left', padx=5)
        self.filter_value = tk.StringVar()
        ttk.Entry(filter_frame, textvariable=self.filter_value).pack(side='left', padx=5)
        
        # Sort options
        sort_frame = ttk.Frame(frame)
        sort_frame.pack(fill='x', padx=20, pady=10)
        
        self.sort_ascending = tk.BooleanVar(value=True)
        ttk.Radiobutton(sort_frame, text="Ascending", variable=self.sort_ascending, value=True).pack(side='left', padx=5)
        ttk.Radiobutton(sort_frame, text="Descending", variable=self.sort_ascending, value=False).pack(side='left', padx=5)
        
        # Action buttons
        button_frame = ttk.Frame(frame)
        button_frame.pack(fill='x', padx=20, pady=10)
        
        ttk.Button(button_frame, text="Apply Filter", command=self.apply_filter).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Sort Data", command=self.sort_data).pack(side='left', padx=5)
        ttk.Button(button_frame, text="Reset", command=self.reset_filter_sort).pack(side='left', padx=5)

    def create_pivot_tab(self, frame):
        ttk.Label(frame, text="Create Pivot Table", style='Header.TLabel').pack(pady=10)
        
        # Column selections
        options_frame = ttk.Frame(frame)
        options_frame.pack(fill='x', padx=20, pady=10)
        
        # Index selection
        ttk.Label(options_frame, text="Index Column:").grid(row=0, column=0, padx=5, pady=5)
        self.pivot_index_var = tk.StringVar()
        self.pivot_index_combo = ttk.Combobox(options_frame, textvariable=self.pivot_index_var)
        self.pivot_index_combo.grid(row=0, column=1, padx=5, pady=5)
        
        # Values selection
        ttk.Label(options_frame, text="Values Column:").grid(row=1, column=0, padx=5, pady=5)
        self.pivot_values_var = tk.StringVar()
        self.pivot_values_combo = ttk.Combobox(options_frame, textvariable=self.pivot_values_var)
        self.pivot_values_combo.grid(row=1, column=1, padx=5, pady=5)
        
        # Aggregation function
        ttk.Label(options_frame, text="Aggregation:").grid(row=2, column=0, padx=5, pady=5)
        self.agg_func_var = tk.StringVar(value='sum')
        agg_funcs = ['sum', 'mean', 'count', 'min', 'max']
        ttk.Combobox(options_frame, textvariable=self.agg_func_var, values=agg_funcs).grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Button(frame, text="Create Pivot Table", command=self.create_pivot_table).pack(pady=10)

    def create_visualization_tab(self, frame):
        ttk.Label(frame, text="Data Visualization", style='Header.TLabel').pack(pady=10)
        
        # Control panel
        control_frame = ttk.Frame(frame)
        control_frame.pack(fill='x', padx=20, pady=10)
        
        # Chart type selection
        ttk.Label(control_frame, text="Chart Type:").grid(row=0, column=0, padx=5, pady=5)
        self.chart_type_var = tk.StringVar(value='line')
        chart_types = ['line', 'bar', 'scatter', 'histogram', 'box']
        ttk.Combobox(control_frame, textvariable=self.chart_type_var, values=chart_types).grid(row=0, column=1, padx=5, pady=5)
        
        # X-axis selection
        ttk.Label(control_frame, text="X-Axis:").grid(row=1, column=0, padx=5, pady=5)
        self.x_axis_var = tk.StringVar()
        self.x_axis_combo = ttk.Combobox(control_frame, textvariable=self.x_axis_var)
        self.x_axis_combo.grid(row=1, column=1, padx=5, pady=5)
        
        # Y-axis selection
        ttk.Label(control_frame, text="Y-Axis:").grid(row=2, column=0, padx=5, pady=5)
        self.y_axis_var = tk.StringVar()
        self.y_axis_combo = ttk.Combobox(control_frame, textvariable=self.y_axis_var)
        self.y_axis_combo.grid(row=2, column=1, padx=5, pady=5)
        
        ttk.Button(frame, text="Generate Chart", command=self.generate_chart).pack(pady=10)

    def create_word_tab(self, frame):
        ttk.Label(frame, text="Convert Excel to Word", style='Header.TLabel').pack(pady=10)
        ttk.Button(frame, text="Convert to Word", command=self.convert_to_word).pack(pady=10)

    def browse_file(self):
        file_path = filedialog.askopenfilename(title="Select an Excel file", filetypes=[("Excel files", "*.xlsx *.xls")])
        if file_path:
            self.selected_file = file_path
            self.file_label.config(text=os.path.basename(file_path))
            self.df = pd.read_excel(self.selected_file)
            self.update_combo_boxes()
            self.status_var.set("Loaded file: " + self.selected_file)

    def update_combo_boxes(self):
        if self.df is not None and not self.df.empty:
            columns = list(self.df.columns)
            self.csv_sheet_combo['values'] = columns
            self.filter_column_combo['values'] = columns
            self.pivot_index_combo['values'] = columns
            self.pivot_values_combo['values'] = columns
            self.x_axis_combo['values'] = columns
            self.y_axis_combo['values'] = columns

    def convert_to_csv(self):
        if self.df is None or self.df.empty:
            messagebox.showwarning("Warning", "No data loaded. Please load an Excel file first.")
            return
            
        output_path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])
        if output_path:
            self.df.to_csv(output_path, index=False)
            messagebox.showinfo("Success", f"Converted to CSV and saved at: {output_path}")

    def apply_filter(self):
        if self.df is None or self.df.empty:
            messagebox.showwarning("Warning", "No data loaded. Please load an Excel file first.")
            return
            
        column = self.filter_column_var.get()
        value = self.filter_value.get()
        
        if not column or not value:
            messagebox.showwarning("Warning", "Please select a column and enter a filter value.")
            return
            
        self.df = self.df[self.df[column].astype(str).str.contains(value, na=False)]
        self.update_combo_boxes()
        self.status_var.set("Applied filter. Rows: {}".format(len(self.df)))

    def sort_data(self):
        if self.df is None or self.df.empty:
            messagebox.showwarning("Warning", "No data loaded. Please load an Excel file first.")
            return
            
        column = self.filter_column_var.get()
        if not column:
            messagebox.showwarning("Warning", "Please select a column to sort by.")
            return
            
        ascending = self.sort_ascending.get()
        self.df = self.df.sort_values(by=column, ascending=ascending)
        self.update_combo_boxes()
        self.status_var.set("Sorted data by column: {}".format(column))

    def reset_filter_sort(self):
        if self.selected_file:
            self.df = pd.read_excel(self.selected_file)
            self.update_combo_boxes()
            self.status_var.set("Reset filter and sort.")

    def create_pivot_table(self):
        if self.df is None or self.df.empty:
            messagebox.showwarning("Warning", "No data loaded. Please load an Excel file first.")
            return

        index_col = self.pivot_index_var.get()
        values_col = self.pivot_values_var.get()
        agg_func = self.agg_func_var.get()

        if not index_col or not values_col:
            messagebox.showwarning("Warning", "Please select both Index and Values columns.")
            return

        pivot_table = self.df.pivot_table(index=index_col, values=values_col, aggfunc=agg_func)
        output_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])

        if output_path:
            with pd.ExcelWriter(output_path) as writer:
                pivot_table.to_excel(writer)
            messagebox.showinfo("Success", f"Pivot table saved to {output_path}")

    def generate_chart(self):
        if self.df is None or self.df.empty:
            messagebox.showwarning("Warning", "No data loaded. Please load an Excel file first.")
            return
            
        chart_type = self.chart_type_var.get()
        x_column = self.x_axis_var.get()
        y_column = self.y_axis_var.get()

        if not x_column or not y_column:
            messagebox.showwarning("Warning", "Please select both X and Y axis columns.")
            return
            
        plt.figure(figsize=(10, 6))
        if chart_type == 'line':
            plt.plot(self.df[x_column], self.df[y_column])
        elif chart_type == 'bar':
            plt.bar(self.df[x_column], self.df[y_column])
        elif chart_type == 'scatter':
            plt.scatter(self.df[x_column], self.df[y_column])
        elif chart_type == 'histogram':
            plt.hist(self.df[y_column], bins=10)
        elif chart_type == 'box':
            sns.boxplot(x=self.df[x_column], y=self.df[y_column])
        
        plt.title(f"{chart_type.capitalize()} Chart")
        plt.xlabel(x_column)
        plt.ylabel(y_column)
        plt.grid(True)
        plt.show()

    def convert_to_word(self):
        if self.df is None or self.df.empty:
            messagebox.showwarning("Warning", "No data loaded. Please load an Excel file first.")
            return
            
        output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if output_path:
            doc = Document()
            doc.add_heading('Excel Data', level=1)

            table = doc.add_table(rows=1, cols=len(self.df.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(self.df.columns):
                hdr_cells[i].text = str(col)

            for index, row in self.df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

            doc.save(output_path)
            messagebox.showinfo("Success", f"Converted to Word and saved at: {output_path}")

if __name__ == "__main__":
    app = EnhancedExcelUtilityApp()
    app.mainloop()